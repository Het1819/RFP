import logging
import uuid
from enum import StrEnum
from typing import Annotated, Any

from fastapi import APIRouter, Depends, Form, HTTPException, Request, Response
from fastapi.responses import HTMLResponse, RedirectResponse
from pydantic import BeforeValidator
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.core import security as core_security
from app.core.config import settings
from app.core.csrf import validate_csrf_token
from app.core.database import get_db, get_default_org_and_user
from app.core.security import (
    ReviewerAuthorizationError,
    get_project_for_org,
    get_requirement_for_org,
    require_requirement_reviewer,
)
from app.core.templates import templates
from app.models.audit import AuditEvent
from app.models.comment import RequirementComment
from app.models.extraction import (
    CANDIDATE_STATUS_PROPOSED,
    MAX_REVIEWER_COMMENT_LEN,
    MAX_REVIEWER_EDITED_TEXT_LEN,
    REVIEW_TASK_STATUS_OPEN,
    CandidateReviewTask,
    ExtractionRun,
    RequirementCandidate,
)
from app.models.project import ProposalProject
from app.models.requirement import Requirement
from app.models.user import User
from app.services.candidate_review import (
    DECISION_APPROVE,
    DECISION_EDIT,
    DECISION_REJECT,
    REVIEW_ALREADY_DECIDED,
    REVIEW_CONFLICT,
    REVIEW_NOT_FOUND,
    CandidateReviewError,
    review_requirement_candidate,
)
from app.services.candidate_review import (
    ReviewResult as CandidateReviewResult,
)
from app.services.project_service import log_audit_event

logger = logging.getLogger(__name__)


def get_current_org_and_user(
    request: Request, db: Session
) -> tuple[uuid.UUID, uuid.UUID]:
    try:
        return core_security.get_current_org_and_user(request, db)
    except HTTPException as e:
        if e.status_code == 401 and settings.AUTH_MODE == "dev":
            core_security.check_app_env_auth()
            logger.warning(
                "Development authentication active. Falling back to default user/org."
            )
            return get_default_org_and_user(db)
        raise


class RequirementStatus(StrEnum):
    NOT_STARTED = "NOT_STARTED"
    NEEDS_EVIDENCE = "NEEDS_EVIDENCE"
    DRAFTED = "DRAFTED"
    NEEDS_REVIEW = "NEEDS_REVIEW"
    APPROVED = "APPROVED"
    REJECTED = "REJECTED"
    IN_REVIEW = "IN_REVIEW"
    CHANGES_REQUESTED = "CHANGES_REQUESTED"


class RiskLevel(StrEnum):
    High = "High"
    Medium = "Medium"
    Low = "Low"


class RequirementType(StrEnum):
    Technical = "Technical"
    Compliance = "Compliance"
    Commercial = "Commercial"
    Procedural = "Procedural"


def empty_to_none(v: Any) -> Any:
    if v == "":
        return None
    return v


OptionalRequirementType = Annotated[
    RequirementType | None, BeforeValidator(empty_to_none)
]
OptionalRiskLevel = Annotated[RiskLevel | None, BeforeValidator(empty_to_none)]

router = APIRouter(tags=["compliance"])


# get_requirement_for_org is imported from app.core.security


@router.get("/projects/{project_id}/matrix", response_class=HTMLResponse)
def matrix_view(
    request: Request, project_id: uuid.UUID, db: Session = Depends(get_db)
) -> Any:
    org_id, user_id = get_current_org_and_user(request, db)
    project = get_project_for_org(db, project_id, org_id)

    # 1. Fetch all requirements for the project to calculate filter counts
    all_reqs = db.scalars(
        select(Requirement).where(Requirement.project_id == project_id)
    ).all()

    assigned_to_me_count = sum(1 for r in all_reqs if r.assigned_to_user_id == user_id)
    unassigned_count = sum(1 for r in all_reqs if r.assigned_to_user_id is None)
    needs_evidence_count = sum(1 for r in all_reqs if r.status == "NEEDS_EVIDENCE")
    needs_review_count = sum(1 for r in all_reqs if r.status == "NEEDS_REVIEW")
    changes_requested_count = sum(
        1 for r in all_reqs if r.status == "CHANGES_REQUESTED"
    )
    approved_count = sum(1 for r in all_reqs if r.status == "APPROVED")
    rejected_count = sum(1 for r in all_reqs if r.status == "REJECTED")

    import datetime

    is_overdue = False
    if project.due_date:
        is_overdue = datetime.datetime.now(datetime.UTC) > project.due_date

    # 2. Apply filtering
    filter_param = request.query_params.get("filter")
    query = select(Requirement).where(Requirement.project_id == project_id)

    if filter_param == "assigned_to_me":
        query = query.where(Requirement.assigned_to_user_id == user_id)
    elif filter_param == "unassigned":
        query = query.where(Requirement.assigned_to_user_id.is_(None))
    elif filter_param == "needs_evidence":
        query = query.where(Requirement.status == "NEEDS_EVIDENCE")
    elif filter_param == "needs_review":
        query = query.where(Requirement.status == "NEEDS_REVIEW")
    elif filter_param == "changes_requested":
        query = query.where(Requirement.status == "CHANGES_REQUESTED")
    elif filter_param == "approved":
        query = query.where(Requirement.status == "APPROVED")
    elif filter_param == "rejected":
        query = query.where(Requirement.status == "REJECTED")

    requirements = db.scalars(
        query.order_by(Requirement.source_page.asc(), Requirement.created_at.asc())
    ).all()

    error_msg = request.query_params.get("error")

    return templates.TemplateResponse(
        request=request,
        name="projects/matrix.html",
        context={
            "project": project,
            "requirements": requirements,
            "error_msg": error_msg,
            "total_count": len(all_reqs),
            "assigned_to_me_count": assigned_to_me_count,
            "unassigned_count": unassigned_count,
            "needs_evidence_count": needs_evidence_count,
            "needs_review_count": needs_review_count,
            "changes_requested_count": changes_requested_count,
            "approved_count": approved_count,
            "rejected_count": rejected_count,
            "is_overdue": is_overdue,
            "current_filter": filter_param or "",
        },
    )


@router.get("/requirements/{requirement_id}/edit", response_class=HTMLResponse)
def edit_requirement_row(
    request: Request, requirement_id: uuid.UUID, db: Session = Depends(get_db)
) -> Any:
    org_id, _ = get_current_org_and_user(request, db)
    req = get_requirement_for_org(db, requirement_id, org_id)
    return templates.TemplateResponse(
        request=request,
        name="projects/matrix_row_edit.html",
        context={"req": req},
    )


@router.post(
    "/requirements/{requirement_id}/edit",
    response_class=HTMLResponse,
    dependencies=[Depends(validate_csrf_token)],
)
def update_requirement_action(
    request: Request,
    requirement_id: uuid.UUID,
    original_text: str = Form(...),
    source_section: str = Form(None),
    source_page: int = Form(None),
    requirement_type: OptionalRequirementType = Form(None),
    mandatory: bool = Form(False),
    status: RequirementStatus = Form(RequirementStatus.NOT_STARTED),
    owner_name: str = Form(None),
    proposal_section: str = Form(None),
    risk_level: OptionalRiskLevel = Form(None),
    db: Session = Depends(get_db),
) -> Any:
    org_id, user_id = get_current_org_and_user(request, db)
    req = get_requirement_for_org(db, requirement_id, org_id)

    old_details = {
        "original_text": req.original_text,
        "status": req.status,
        "mandatory": req.mandatory,
    }

    req.original_text = original_text
    req.source_section = source_section
    req.source_page = source_page
    req.requirement_type = requirement_type
    req.mandatory = mandatory
    req.status = status
    req.owner_name = owner_name
    req.proposal_section = proposal_section
    req.risk_level = risk_level
    db.commit()

    log_audit_event(
        db,
        org_id=org_id,
        user_id=user_id,
        action="requirement_edit",
        entity_type="Requirement",
        entity_id=req.id,
        details={
            "old": old_details,
            "new": {
                "original_text": original_text,
                "status": status,
                "mandatory": mandatory,
            },
        },
    )

    return templates.TemplateResponse(
        request=request,
        name="projects/matrix_row.html",
        context={"req": req},
    )


@router.get("/requirements/{requirement_id}/cancel", response_class=HTMLResponse)
def cancel_edit_row(
    request: Request, requirement_id: uuid.UUID, db: Session = Depends(get_db)
) -> Any:
    org_id, _ = get_current_org_and_user(request, db)
    req = get_requirement_for_org(db, requirement_id, org_id)
    return templates.TemplateResponse(
        request=request,
        name="projects/matrix_row.html",
        context={"req": req},
    )


@router.delete(
    "/requirements/{requirement_id}",
    response_class=HTMLResponse,
    dependencies=[Depends(validate_csrf_token)],
)
def delete_requirement_action(
    request: Request, requirement_id: uuid.UUID, db: Session = Depends(get_db)
) -> Any:
    org_id, user_id = get_current_org_and_user(request, db)
    req = get_requirement_for_org(db, requirement_id, org_id)
    db.delete(req)
    db.commit()

    log_audit_event(
        db,
        org_id=org_id,
        user_id=user_id,
        action="requirement_delete",
        entity_type="Requirement",
        entity_id=requirement_id,
        details={"original_text": req.original_text},
    )

    return HTMLResponse(content="")


@router.post(
    "/projects/{project_id}/matrix/merge",
    response_class=RedirectResponse,
    dependencies=[Depends(validate_csrf_token)],
)
def merge_requirements_action(
    project_id: uuid.UUID,
    request: Request,
    ids: list[str] = Form(...),
    db: Session = Depends(get_db),
) -> Any:
    if not ids or len(ids) < 2:
        url = f"/projects/{project_id}/matrix?error=Select at least two items to merge"
        return RedirectResponse(
            url=url,
            status_code=303,
        )

    org_id, user_id = get_current_org_and_user(request, db)
    project = get_project_for_org(db, project_id, org_id)
    req_ids = [uuid.UUID(i) for i in ids]

    reqs = db.scalars(
        select(Requirement).where(
            Requirement.id.in_(req_ids),
            Requirement.project_id == project.id,
        )
    ).all()

    if len(reqs) != len(req_ids):
        raise HTTPException(status_code=404, detail="No requirements found")

    merged_text = "\n[Merged] ".join([r.original_text for r in reqs])
    primary = reqs[0]

    primary.original_text = merged_text
    db.commit()

    for secondary in reqs[1:]:
        db.delete(secondary)
    db.commit()

    log_audit_event(
        db,
        org_id=org_id,
        user_id=user_id,
        action="requirements_merge",
        entity_type="Requirement",
        entity_id=primary.id,
        details={"merged_ids": [str(r.id) for r in reqs]},
    )

    return RedirectResponse(url=f"/projects/{project_id}/matrix", status_code=303)


@router.post(
    "/requirements/{requirement_id}/split",
    response_class=RedirectResponse,
    dependencies=[Depends(validate_csrf_token)],
)
def split_requirement_action(
    requirement_id: uuid.UUID,
    request: Request,
    split_text: str = Form(...),
    db: Session = Depends(get_db),
) -> Any:
    org_id, user_id = get_current_org_and_user(request, db)
    req = get_requirement_for_org(db, requirement_id, org_id)

    secondary = Requirement(
        project_id=req.project_id,
        source_document_id=req.source_document_id,
        original_text=split_text,
        source_section=req.source_section,
        source_page=req.source_page,
        requirement_type=req.requirement_type,
        mandatory=req.mandatory,
        status="NOT_STARTED",
        risk_level=req.risk_level,
    )
    db.add(secondary)

    req.original_text = req.original_text.replace(split_text, "").strip()
    if not req.original_text:
        req.original_text = "[Split Part 1]"

    db.commit()

    log_audit_event(
        db,
        org_id=org_id,
        user_id=user_id,
        action="requirement_split",
        entity_type="Requirement",
        entity_id=req.id,
        details={"new_requirement_id": str(secondary.id)},
    )

    return RedirectResponse(url=f"/projects/{req.project_id}/matrix", status_code=303)


@router.get("/requirements/{requirement_id}/workspace", response_class=HTMLResponse)
def requirement_workspace_view(
    request: Request, requirement_id: uuid.UUID, db: Session = Depends(get_db)
) -> Any:
    org_id, user_id = get_current_org_and_user(request, db)
    req = get_requirement_for_org(db, requirement_id, org_id)
    project = get_project_for_org(db, req.project_id, org_id)

    # Run retrieval
    from app.services.retriever import retrieve_evidence

    q_param = request.query_params.get("q")
    if q_param is not None:
        search_query = q_param
    else:
        search_query = " ".join(req.original_text.split()[:8])
    evidence_passages = retrieve_evidence(db, req.project_id, search_query)

    # Get existing linked evidence
    from app.models.evidence import EvidenceLink

    linked_evidence = db.scalars(
        select(EvidenceLink).where(EvidenceLink.requirement_id == requirement_id)
    ).all()

    # Get draft response
    from app.models.response import DraftResponse

    draft = db.scalars(
        select(DraftResponse)
        .where(DraftResponse.requirement_id == requirement_id)
        .order_by(DraftResponse.version.desc())
    ).first()

    # Get org users for reviewer assignment dropdown
    users = db.scalars(
        select(User)
        .where(User.organization_id == org_id)
        .order_by(User.full_name.asc())
    ).all()

    # Get requirement comments
    comments = db.scalars(
        select(RequirementComment)
        .where(RequirementComment.requirement_id == requirement_id)
        .order_by(RequirementComment.created_at.desc())
    ).all()

    # Get recent audit events
    recent_audits = db.scalars(
        select(AuditEvent)
        .where(AuditEvent.entity_id == requirement_id)
        .order_by(AuditEvent.created_at.desc())
        .limit(10)
    ).all()

    warning = request.query_params.get("warning")

    return templates.TemplateResponse(
        request=request,
        name="projects/requirement_workspace.html",
        context={
            "project": project,
            "req": req,
            "evidence_passages": evidence_passages,
            "linked_evidence": linked_evidence,
            "draft": draft,
            "search_query": search_query,
            "users": users,
            "comments": comments,
            "recent_audits": recent_audits,
            "warning": warning,
            "current_user_id": user_id,
        },
    )


@router.post(
    "/requirements/{requirement_id}/evidence/link",
    response_class=RedirectResponse,
    dependencies=[Depends(validate_csrf_token)],
)
def link_evidence_action(
    requirement_id: uuid.UUID,
    request: Request,
    document_id: uuid.UUID = Form(...),
    snippet: str = Form(...),
    page_number: int = Form(None),
    # client score is ignored — clamped/recomputed server-side
    score: float = Form(0.0),
    db: Session = Depends(get_db),
) -> Any:
    org_id, user_id = get_current_org_and_user(request, db)
    req = get_requirement_for_org(db, requirement_id, org_id)

    # Validate evidence candidate server-side — never trust client snippet/page/score
    from app.services.evidence_validation import (
        EvidenceValidationError,
        validate_evidence_candidate,
    )

    try:
        canonical_snippet, resolved_page = validate_evidence_candidate(
            db,
            requirement_project_id=req.project_id,
            document_id=document_id,
            page_number=page_number,
            client_snippet=snippet,
        )
    except EvidenceValidationError as exc:
        from app.core.observability import MetricsRegistry

        MetricsRegistry.evidence_validation_failures += 1
        logger.warning(f"Evidence validation failed: {exc.detail}")
        raise HTTPException(status_code=exc.status_code, detail=exc.detail) from exc

    # Client-submitted score is NEVER stored; default 0.0 (retriever sets real scores)
    _ = score  # explicitly discard
    server_score = 0.0  # TODO: compute from FTS rank if available

    from app.models.evidence import EvidenceLink

    link = EvidenceLink(
        requirement_id=requirement_id,
        document_id=document_id,
        snippet=canonical_snippet,  # server-resolved canonical snippet
        page_number=resolved_page,
        score=server_score,
    )
    db.add(link)
    db.commit()

    log_audit_event(
        db,
        org_id=org_id,
        user_id=user_id,
        action="evidence_link",
        entity_type="EvidenceLink",
        entity_id=link.id,
        details={
            "requirement_id": str(requirement_id),
            "document_id": str(document_id),
            "page_number": resolved_page,
            "snippet_len": len(canonical_snippet),
        },
    )

    return RedirectResponse(
        url=f"/requirements/{requirement_id}/workspace", status_code=303
    )


@router.post(
    "/requirements/{requirement_id}/draft",
    response_class=RedirectResponse,
    dependencies=[Depends(validate_csrf_token)],
)
async def draft_requirement_response(
    requirement_id: uuid.UUID,
    request: Request,
    db: Session = Depends(get_db),
) -> Any:
    org_id, user_id = get_current_org_and_user(request, db)
    req = get_requirement_for_org(db, requirement_id, org_id)

    from app.models.evidence import EvidenceLink

    links = db.scalars(
        select(EvidenceLink).where(EvidenceLink.requirement_id == requirement_id)
    ).all()

    evidence_snippets = [
        {
            "doc_id": str(link.document_id),
            "snippet": link.snippet,
            "page_number": link.page_number,
            "score": link.score,
        }
        for link in links
    ]

    from app.core.llm import get_llm_provider

    provider = get_llm_provider()

    draft_draft = await provider.draft_response(req.original_text, evidence_snippets)

    from sqlalchemy import func

    from app.models.response import DraftResponse

    max_version = db.scalar(
        select(func.max(DraftResponse.version)).where(
            DraftResponse.requirement_id == requirement_id
        )
    )
    new_version = (max_version or 0) + 1

    draft = DraftResponse(
        requirement_id=requirement_id,
        content=draft_draft.answer_text,
        confidence=draft_draft.confidence,
        needs_evidence=draft_draft.needs_evidence,
        assumptions=draft_draft.assumptions,
        status="draft",
        version=new_version,
    )
    db.add(draft)

    if draft_draft.needs_evidence:
        req.status = "NEEDS_EVIDENCE"
    else:
        req.status = "DRAFTED"

    db.commit()

    log_audit_event(
        db,
        org_id=org_id,
        user_id=user_id,
        action="draft_generate",
        entity_type="DraftResponse",
        entity_id=draft.id,
        details={
            "requirement_id": str(requirement_id),
            "needs_evidence": draft.needs_evidence,
        },
    )

    return RedirectResponse(
        url=f"/requirements/{requirement_id}/workspace", status_code=303
    )


@router.post(
    "/requirements/{requirement_id}/draft/edit",
    response_class=RedirectResponse,
    dependencies=[Depends(validate_csrf_token)],
)
def edit_draft_response(
    requirement_id: uuid.UUID,
    request: Request,
    content: str = Form(...),
    db: Session = Depends(get_db),
) -> Any:
    org_id, user_id = get_current_org_and_user(request, db)
    req = get_requirement_for_org(db, requirement_id, org_id)

    from app.models.response import DraftResponse

    draft = db.scalars(
        select(DraftResponse)
        .where(DraftResponse.requirement_id == requirement_id)
        .order_by(DraftResponse.version.desc())
    ).first()

    if not draft:
        raise HTTPException(status_code=404, detail="Draft not found")

    draft.content = content
    req.status = "DRAFTED"
    db.commit()

    log_audit_event(
        db,
        org_id=org_id,
        user_id=user_id,
        action="draft_edit",
        entity_type="DraftResponse",
        entity_id=draft.id,
        details={"requirement_id": str(requirement_id)},
    )

    return RedirectResponse(
        url=f"/requirements/{requirement_id}/workspace", status_code=303
    )


@router.post(
    "/requirements/{requirement_id}/draft/approve",
    response_class=RedirectResponse,
    dependencies=[Depends(validate_csrf_token)],
)
def approve_draft_response(
    requirement_id: uuid.UUID,
    request: Request,
    db: Session = Depends(get_db),
) -> Any:
    org_id, user_id = get_current_org_and_user(request, db)
    req = get_requirement_for_org(db, requirement_id, org_id)

    from app.models.evidence import EvidenceLink
    from app.models.response import DraftResponse
    from app.services.evidence_validation import validate_draft_grounding

    draft = db.scalars(
        select(DraftResponse)
        .where(DraftResponse.requirement_id == requirement_id)
        .order_by(DraftResponse.version.desc())
    ).first()

    if not draft:
        raise HTTPException(status_code=404, detail="Draft not found")

    # Check grounding: collect validated evidence snippets
    links = db.scalars(
        select(EvidenceLink).where(EvidenceLink.requirement_id == requirement_id)
    ).all()

    # Block approval of mandatory requirements that have zero evidence
    if req.mandatory and not links:
        draft.status = "needs_review"
        req.status = "NEEDS_REVIEW"
        db.commit()
        log_audit_event(
            db,
            org_id=org_id,
            user_id=user_id,
            action="draft_approve_blocked",
            entity_type="DraftResponse",
            entity_id=draft.id,
            details={
                "requirement_id": str(requirement_id),
                "reason": "mandatory_requirement_no_evidence",
            },
        )
        return RedirectResponse(
            url=(
                f"/requirements/{requirement_id}/workspace"
                "?warning=mandatory_requirement_needs_evidence"
            ),
            status_code=303,
        )

    # Run grounding check against validated evidence snippets
    evidence_snippets = [link.snippet for link in links]
    grounding = validate_draft_grounding(draft.content, evidence_snippets)

    if not grounding.passes and evidence_snippets:
        from app.core.observability import MetricsRegistry

        MetricsRegistry.evidence_validation_failures += 1
        # Draft has unsupported claims — route to NEEDS_REVIEW instead of APPROVED
        draft.status = "needs_review"
        req.status = "NEEDS_REVIEW"
        db.commit()
        log_audit_event(
            db,
            org_id=org_id,
            user_id=user_id,
            action="draft_approve_blocked",
            entity_type="DraftResponse",
            entity_id=draft.id,
            details={
                "requirement_id": str(requirement_id),
                "reason": "unsupported_claims",
                "unsupported_count": len(grounding.unsupported_claims),
                "grounding_pass_rate": grounding.grounding_pass_rate,
            },
        )
        return RedirectResponse(
            url=(
                f"/requirements/{requirement_id}/workspace?warning=unsupported_claims"
            ),
            status_code=303,
        )

    draft.status = "approved"
    draft.approved_by_id = user_id
    req.status = "APPROVED"
    db.commit()

    log_audit_event(
        db,
        org_id=org_id,
        user_id=user_id,
        action="draft_approve",
        entity_type="DraftResponse",
        entity_id=draft.id,
        details={"requirement_id": str(requirement_id)},
    )

    log_audit_event(
        db,
        org_id=org_id,
        user_id=user_id,
        action="REVIEW_APPROVED",
        entity_type="Requirement",
        entity_id=req.id,
        details={"draft_response_id": str(draft.id)},
    )

    return RedirectResponse(
        url=f"/requirements/{requirement_id}/workspace", status_code=303
    )


@router.post(
    "/requirements/{requirement_id}/draft/reject",
    response_class=RedirectResponse,
    dependencies=[Depends(validate_csrf_token)],
)
def reject_draft_response(
    requirement_id: uuid.UUID,
    request: Request,
    reason: str = Form(None),
    db: Session = Depends(get_db),
) -> Any:
    org_id, user_id = get_current_org_and_user(request, db)
    req = get_requirement_for_org(db, requirement_id, org_id)

    from app.models.response import DraftResponse

    draft = db.scalars(
        select(DraftResponse)
        .where(DraftResponse.requirement_id == requirement_id)
        .order_by(DraftResponse.version.desc())
    ).first()

    if not draft:
        raise HTTPException(status_code=404, detail="Draft not found")

    draft.status = "rejected"
    req.status = "REJECTED"

    reason_text = reason.strip() if reason else ""
    if reason_text:
        from app.models.comment import RequirementComment

        comment = RequirementComment(
            requirement_id=requirement_id,
            author_user_id=user_id,
            content=reason_text,
            decision_type="REJECTED",
        )
        db.add(comment)

    db.commit()

    log_audit_event(
        db,
        org_id=org_id,
        user_id=user_id,
        action="draft_reject",
        entity_type="DraftResponse",
        entity_id=draft.id,
        details={"requirement_id": str(requirement_id)},
    )

    log_audit_event(
        db,
        org_id=org_id,
        user_id=user_id,
        action="REVIEW_REJECTED",
        entity_type="Requirement",
        entity_id=req.id,
        details={"reason": reason_text or "No reason provided"},
    )

    return RedirectResponse(
        url=f"/requirements/{requirement_id}/workspace", status_code=303
    )


@router.post(
    "/requirements/{requirement_id}/draft/regenerate",
    response_class=RedirectResponse,
    dependencies=[Depends(validate_csrf_token)],
)
async def regenerate_draft_response(
    requirement_id: uuid.UUID,
    request: Request,
    db: Session = Depends(get_db),
) -> Any:
    return await draft_requirement_response(requirement_id, request, db)


@router.post(
    "/requirements/{requirement_id}/assign",
    response_class=RedirectResponse,
    dependencies=[Depends(validate_csrf_token)],
)
def assign_requirement_reviewer(
    requirement_id: uuid.UUID,
    request: Request,
    assigned_to_user_id: str = Form(None),
    reviewer_name: str = Form(None),
    db: Session = Depends(get_db),
) -> Any:
    org_id, user_id = get_current_org_and_user(request, db)
    req = get_requirement_for_org(db, requirement_id, org_id)

    target_user_id = None
    target_name = None

    if assigned_to_user_id and assigned_to_user_id != "unassign":
        try:
            target_uuid = uuid.UUID(assigned_to_user_id)
        except ValueError:
            raise HTTPException(
                status_code=400, detail="Invalid user ID format"
            ) from None

        target_user = db.scalar(
            select(User).where(User.id == target_uuid, User.organization_id == org_id)
        )
        if not target_user:
            raise HTTPException(status_code=404, detail="User not found")

        target_user_id = target_user.id
        target_name = target_user.full_name
    elif reviewer_name:
        # Fallback for backward compatibility tests
        target_name = reviewer_name

    req.assigned_to_user_id = target_user_id
    req.assigned_by_user_id = user_id
    import datetime

    req.assigned_at = datetime.datetime.now(datetime.UTC) if target_user_id else None
    req.owner_name = target_name

    # Set status to NEEDS_REVIEW when assigned
    req.status = "NEEDS_REVIEW"

    from app.models.review import ReviewTask

    task = ReviewTask(
        requirement_id=requirement_id,
        assigned_to_id=target_user_id,
        reviewer_notes=f"Routed to {target_name or 'unassigned'} for review.",
        status="open",
    )
    db.add(task)
    db.commit()

    log_audit_event(
        db,
        org_id=org_id,
        user_id=user_id,
        action="REVIEW_ASSIGNED",
        entity_type="Requirement",
        entity_id=requirement_id,
        details={
            "reviewer_name": target_name,
            "reviewer_user_id": str(target_user_id) if target_user_id else None,
            "review_task_id": str(task.id),
        },
    )

    # Maintain old audit action for backward compatibility tests
    log_audit_event(
        db,
        org_id=org_id,
        user_id=user_id,
        action="requirement_assign",
        entity_type="Requirement",
        entity_id=requirement_id,
        details={
            "reviewer_name": target_name,
            "review_task_id": str(task.id),
        },
    )

    return RedirectResponse(
        url=f"/requirements/{requirement_id}/workspace", status_code=303
    )


@router.post(
    "/requirements/{requirement_id}/review/start",
    response_class=RedirectResponse,
    dependencies=[Depends(validate_csrf_token)],
)
def start_review_action(
    requirement_id: uuid.UUID,
    request: Request,
    db: Session = Depends(get_db),
) -> Any:
    org_id, user_id = get_current_org_and_user(request, db)
    req = get_requirement_for_org(db, requirement_id, org_id)

    req.status = "IN_REVIEW"
    db.commit()

    log_audit_event(
        db,
        org_id=org_id,
        user_id=user_id,
        action="REVIEW_STARTED",
        entity_type="Requirement",
        entity_id=req.id,
        details={"status": "IN_REVIEW"},
    )
    return RedirectResponse(
        url=f"/requirements/{requirement_id}/workspace", status_code=303
    )


@router.post(
    "/requirements/{requirement_id}/review/changes-requested",
    response_class=RedirectResponse,
    dependencies=[Depends(validate_csrf_token)],
)
def request_changes_action(
    requirement_id: uuid.UUID,
    request: Request,
    reason: str = Form(...),
    db: Session = Depends(get_db),
) -> Any:
    org_id, user_id = get_current_org_and_user(request, db)
    req = get_requirement_for_org(db, requirement_id, org_id)

    reason_stripped = reason.strip()
    if not reason_stripped:
        raise HTTPException(status_code=400, detail="Reason cannot be empty")

    if len(reason_stripped) > 4000:
        raise HTTPException(status_code=400, detail="Reason is too long")

    req.status = "CHANGES_REQUESTED"

    from app.models.response import DraftResponse

    draft = db.scalars(
        select(DraftResponse)
        .where(DraftResponse.requirement_id == requirement_id)
        .order_by(DraftResponse.version.desc())
    ).first()
    if draft:
        draft.status = "changes_requested"

    from app.models.comment import RequirementComment

    comment = RequirementComment(
        requirement_id=requirement_id,
        author_user_id=user_id,
        content=reason_stripped,
        decision_type="CHANGES_REQUESTED",
    )
    db.add(comment)
    db.commit()

    log_audit_event(
        db,
        org_id=org_id,
        user_id=user_id,
        action="REVIEW_CHANGES_REQUESTED",
        entity_type="Requirement",
        entity_id=req.id,
        details={"reason": reason_stripped[:200]},
    )
    return RedirectResponse(
        url=f"/requirements/{requirement_id}/workspace", status_code=303
    )


@router.post(
    "/requirements/{requirement_id}/review/reject",
    response_class=RedirectResponse,
    dependencies=[Depends(validate_csrf_token)],
)
def reject_review_action(
    requirement_id: uuid.UUID,
    request: Request,
    reason: str = Form(...),
    db: Session = Depends(get_db),
) -> Any:
    org_id, user_id = get_current_org_and_user(request, db)
    req = get_requirement_for_org(db, requirement_id, org_id)

    reason_stripped = reason.strip()
    if not reason_stripped:
        raise HTTPException(status_code=400, detail="Reason cannot be empty")

    if len(reason_stripped) > 4000:
        raise HTTPException(status_code=400, detail="Reason is too long")

    req.status = "REJECTED"

    from app.models.response import DraftResponse

    draft = db.scalars(
        select(DraftResponse)
        .where(DraftResponse.requirement_id == requirement_id)
        .order_by(DraftResponse.version.desc())
    ).first()
    if draft:
        draft.status = "rejected"

    from app.models.comment import RequirementComment

    comment = RequirementComment(
        requirement_id=requirement_id,
        author_user_id=user_id,
        content=reason_stripped,
        decision_type="REJECTED",
    )
    db.add(comment)
    db.commit()

    log_audit_event(
        db,
        org_id=org_id,
        user_id=user_id,
        action="REVIEW_REJECTED",
        entity_type="Requirement",
        entity_id=req.id,
        details={"reason": reason_stripped[:200]},
    )
    return RedirectResponse(
        url=f"/requirements/{requirement_id}/workspace", status_code=303
    )


@router.post(
    "/requirements/{requirement_id}/review/reopen",
    response_class=RedirectResponse,
    dependencies=[Depends(validate_csrf_token)],
)
def reopen_review_action(
    requirement_id: uuid.UUID,
    request: Request,
    db: Session = Depends(get_db),
) -> Any:
    org_id, user_id = get_current_org_and_user(request, db)
    req = get_requirement_for_org(db, requirement_id, org_id)

    req.status = "NEEDS_REVIEW"

    from app.models.response import DraftResponse

    draft = db.scalars(
        select(DraftResponse)
        .where(DraftResponse.requirement_id == requirement_id)
        .order_by(DraftResponse.version.desc())
    ).first()
    if draft:
        draft.status = "draft"

    db.commit()

    log_audit_event(
        db,
        org_id=org_id,
        user_id=user_id,
        action="REVIEW_REOPENED",
        entity_type="Requirement",
        entity_id=req.id,
        details={"status": "NEEDS_REVIEW"},
    )
    return RedirectResponse(
        url=f"/requirements/{requirement_id}/workspace", status_code=303
    )


@router.post(
    "/requirements/{requirement_id}/comments",
    response_class=RedirectResponse,
    dependencies=[Depends(validate_csrf_token)],
)
def add_comment_action(
    requirement_id: uuid.UUID,
    request: Request,
    content: str = Form(...),
    db: Session = Depends(get_db),
) -> Any:
    org_id, user_id = get_current_org_and_user(request, db)
    req = get_requirement_for_org(db, requirement_id, org_id)

    content_stripped = content.strip()
    if not content_stripped:
        raise HTTPException(status_code=400, detail="Comment cannot be empty")

    if len(content_stripped) > 4000:
        raise HTTPException(status_code=400, detail="Comment is too long")

    from app.models.comment import RequirementComment

    comment = RequirementComment(
        requirement_id=requirement_id,
        author_user_id=user_id,
        content=content_stripped,
        decision_type="NOTE",
    )
    db.add(comment)
    db.commit()

    log_audit_event(
        db,
        org_id=org_id,
        user_id=user_id,
        action="REVIEW_NOTE_ADDED",
        entity_type="Requirement",
        entity_id=req.id,
        details={"comment_id": str(comment.id)},
    )
    return RedirectResponse(
        url=f"/requirements/{requirement_id}/workspace", status_code=303
    )


# ---------------------------------------------------------------------------
# Requirement candidate review (A5f Pass 2A)
# ---------------------------------------------------------------------------
# Form routes only -- the application has no JSON API surface, and adding one
# for this would widen the attack surface for no product gain.
#
# Every route is deliberately thin: it resolves the session identity, hands the
# candidate id and the decision to the service, and translates the service's
# fixed result code into a response. Authorization, locking, source
# revalidation, and the Requirement insert all live in the service, so the
# authority boundary cannot be bypassed by any future non-HTTP caller.
#
# Nothing about the candidate is accepted from the client except its id and the
# reviewer's own text: no status, reviewer id, run id, organization id, project
# id, provenance, or source-candidate linkage is bindable from a form field.


# Bounded rendering limits. Evidence is quoted from an untrusted document, so
# the queue shows an excerpt rather than an unbounded block; the detail page
# shows the full slice a reviewer must actually read to make the decision.
CANDIDATE_QUEUE_PAGE_SIZE = 25
CANDIDATE_EVIDENCE_EXCERPT_CHARS = 300


def _require_reviewer(request: Request, db: Session) -> tuple[uuid.UUID, User]:
    """Resolve the session identity and enforce the reviewer capability.

    Runs before any candidate is read, so a caller without the capability
    cannot use these pages to learn whether candidates exist -- not their
    count, not their projects, not their ids.
    """
    org_id, user_id = get_current_org_and_user(request, db)
    reviewer = require_requirement_reviewer(db, user_id, org_id)
    return org_id, reviewer


def _open_candidate_query(org_id: uuid.UUID) -> Any:
    """Open review work for one organization, oldest first.

    Joined against CandidateReviewTask rather than filtered on candidate status
    alone: a task that was completed or superseded is not open work, even if
    something later reopened the candidate. Both sides are tenant-filtered so a
    mismatch cannot widen the result set.
    """
    return (
        select(RequirementCandidate, CandidateReviewTask)
        .join(
            CandidateReviewTask,
            CandidateReviewTask.candidate_id == RequirementCandidate.id,
        )
        .where(
            RequirementCandidate.organization_id == org_id,
            CandidateReviewTask.organization_id == org_id,
            RequirementCandidate.candidate_status == CANDIDATE_STATUS_PROPOSED,
            CandidateReviewTask.status == REVIEW_TASK_STATUS_OPEN,
        )
        .order_by(RequirementCandidate.created_at.asc(), RequirementCandidate.id.asc())
    )


def _excerpt(text: str, limit: int = CANDIDATE_EVIDENCE_EXCERPT_CHARS) -> str:
    collapsed = " ".join((text or "").split())
    if len(collapsed) <= limit:
        return collapsed
    return collapsed[:limit].rstrip() + "…"


@router.get("/compliance/requirement-candidates", response_class=HTMLResponse)
def requirement_candidate_queue(
    request: Request,
    page: int = 1,
    project_id: str | None = None,
    db: Session = Depends(get_db),
) -> Any:
    """Reviewer queue: open candidate review tasks for this organization."""
    org_id, reviewer = _require_reviewer(request, db)

    query = _open_candidate_query(org_id)

    # Optional project filter, validated against the caller's organization so a
    # foreign project id filters to nothing rather than leaking its existence.
    selected_project: uuid.UUID | None = None
    if project_id:
        try:
            candidate_project = uuid.UUID(project_id)
        except ValueError:
            raise HTTPException(status_code=400, detail="Invalid project id") from None
        owned = db.scalar(
            select(ProposalProject.id).where(
                ProposalProject.id == candidate_project,
                ProposalProject.organization_id == org_id,
            )
        )
        if owned is None:
            raise HTTPException(status_code=404, detail="Not found")
        selected_project = owned
        query = query.where(RequirementCandidate.project_id == selected_project)

    page = max(1, page)
    offset = (page - 1) * CANDIDATE_QUEUE_PAGE_SIZE
    # Fetch one extra row to detect a next page without a second count query,
    # which would also have to be tenant-scoped to avoid leaking totals.
    rows = list(db.execute(query.offset(offset).limit(CANDIDATE_QUEUE_PAGE_SIZE + 1)))
    has_next = len(rows) > CANDIDATE_QUEUE_PAGE_SIZE
    rows = rows[:CANDIDATE_QUEUE_PAGE_SIZE]

    project_names = {
        row.id: row.name
        for row in db.execute(
            select(ProposalProject.id, ProposalProject.name).where(
                ProposalProject.organization_id == org_id
            )
        )
    }

    items = [
        {
            "candidate_id": candidate.id,
            "project_id": candidate.project_id,
            "project_name": project_names.get(candidate.project_id, "—"),
            "requirement_text": candidate.normalized_requirement_text,
            "evidence_excerpt": _excerpt(candidate.evidence_text),
            "unit_kind": candidate.unit_kind,
            "source_locator": candidate.source_locator,
            "requirement_type": candidate.requirement_type,
            "confidence": candidate.confidence,
            "created_at": candidate.created_at,
        }
        for candidate, _task in rows
    ]

    projects = sorted(project_names.items(), key=lambda pair: pair[1])

    return templates.TemplateResponse(
        request=request,
        name="compliance/candidate_queue.html",
        context={
            "items": items,
            "page": page,
            "has_next": has_next,
            "has_prev": page > 1,
            "projects": projects,
            "selected_project_id": str(selected_project) if selected_project else "",
            "reviewer_name": reviewer.full_name,
            "notice": request.query_params.get("notice"),
        },
    )


@router.get(
    "/compliance/requirement-candidates/{candidate_id}", response_class=HTMLResponse
)
def requirement_candidate_detail(
    candidate_id: uuid.UUID,
    request: Request,
    db: Session = Depends(get_db),
) -> Any:
    """Candidate detail: immutable evidence plus the three review actions."""
    org_id, _reviewer = _require_reviewer(request, db)

    candidate = db.scalar(
        select(RequirementCandidate).where(
            RequirementCandidate.id == candidate_id,
            RequirementCandidate.organization_id == org_id,
        )
    )
    if candidate is None:
        # Same response for "absent" and "another tenant's".
        raise HTTPException(status_code=404, detail="Not found")

    task = db.scalar(
        select(CandidateReviewTask).where(
            CandidateReviewTask.candidate_id == candidate.id,
            CandidateReviewTask.organization_id == org_id,
        )
    )
    project = db.get(ProposalProject, candidate.project_id)
    run = db.get(ExtractionRun, candidate.extraction_run_id)

    promoted = db.scalar(
        select(Requirement).where(Requirement.source_candidate_id == candidate.id)
    )

    return templates.TemplateResponse(
        request=request,
        name="compliance/candidate_detail.html",
        context={
            "candidate": candidate,
            "task": task,
            "project_name": project.name if project else "—",
            # Version metadata is safe: identifiers, not prompt or model output.
            "prompt_version": run.prompt_version if run else None,
            "schema_version": candidate.extraction_schema_version,
            "provider": run.provider if run else None,
            "model": run.model if run else None,
            "is_open": candidate.candidate_status == CANDIDATE_STATUS_PROPOSED,
            "promoted_requirement_id": promoted.id if promoted else None,
            "max_edit_len": MAX_REVIEWER_EDITED_TEXT_LEN,
            "max_comment_len": MAX_REVIEWER_COMMENT_LEN,
        },
    )


def _candidate_review_response(
    request: Request, result: CandidateReviewResult
) -> Response:
    """Translate a completed review into an HTMX-friendly response."""
    if request.headers.get("hx-request"):
        return HTMLResponse(
            content=(
                f'<div class="candidate-reviewed" '
                f'data-candidate-id="{result.candidate_id}" '
                f'data-status="{result.candidate_status}">Review recorded.</div>'
            ),
            status_code=200,
        )
    # Back to the queue: the reviewed candidate has left it, so the next item
    # is already at the top.
    return RedirectResponse(
        url="/compliance/requirement-candidates?notice=review_recorded",
        status_code=303,
    )


def _review_candidate_action(
    request: Request,
    db: Session,
    candidate_id: uuid.UUID,
    decision: str,
    edited_text: str | None = None,
    reviewer_comment: str | None = None,
) -> Response:
    org_id, user_id = get_current_org_and_user(request, db)
    try:
        result = review_requirement_candidate(
            db,
            candidate_id=candidate_id,
            reviewer_id=user_id,
            org_id=org_id,
            decision=decision,
            edited_text=edited_text,
            reviewer_comment=reviewer_comment,
        )
    except ReviewerAuthorizationError:
        # Already a non-disclosing HTTPException (404 cross-tenant, 403 for a
        # member without the capability). Re-raise unchanged.
        raise
    except CandidateReviewError as err:
        if err.code == REVIEW_NOT_FOUND:
            raise HTTPException(status_code=404, detail="Not found") from None
        if err.code == REVIEW_CONFLICT or err.code == REVIEW_ALREADY_DECIDED:
            raise HTTPException(
                status_code=409, detail="Candidate has already been reviewed"
            ) from None
        # Fixed, non-leaking message for every remaining failure (source drift,
        # bad reviewer text, missing task).
        raise HTTPException(status_code=400, detail=err.code) from None

    return _candidate_review_response(request, result)


@router.post(
    "/compliance/requirement-candidates/{candidate_id}/approve",
    dependencies=[Depends(validate_csrf_token)],
)
def approve_requirement_candidate(
    candidate_id: uuid.UUID,
    request: Request,
    reviewer_comment: str = Form(None),
    db: Session = Depends(get_db),
) -> Response:
    return _review_candidate_action(
        request,
        db,
        candidate_id,
        DECISION_APPROVE,
        reviewer_comment=reviewer_comment,
    )


@router.post(
    "/compliance/requirement-candidates/{candidate_id}/edit",
    dependencies=[Depends(validate_csrf_token)],
)
def edit_requirement_candidate(
    candidate_id: uuid.UUID,
    request: Request,
    edited_text: str = Form(...),
    reviewer_comment: str = Form(None),
    db: Session = Depends(get_db),
) -> Response:
    return _review_candidate_action(
        request,
        db,
        candidate_id,
        DECISION_EDIT,
        edited_text=edited_text,
        reviewer_comment=reviewer_comment,
    )


@router.post(
    "/compliance/requirement-candidates/{candidate_id}/reject",
    dependencies=[Depends(validate_csrf_token)],
)
def reject_requirement_candidate(
    candidate_id: uuid.UUID,
    request: Request,
    reviewer_comment: str = Form(None),
    db: Session = Depends(get_db),
) -> Response:
    return _review_candidate_action(
        request,
        db,
        candidate_id,
        DECISION_REJECT,
        reviewer_comment=reviewer_comment,
    )


@router.get("/projects/{project_id}/export/matrix")
def export_compliance_matrix(
    project_id: uuid.UUID,
    request: Request,
    db: Session = Depends(get_db),
) -> Any:
    org_id, user_id = get_current_org_and_user(request, db)
    _ = get_project_for_org(db, project_id, org_id)

    log_audit_event(
        db,
        org_id=org_id,
        user_id=user_id,
        action="EXPORT_COMPLIANCE_MATRIX",
        entity_type="Project",
        entity_id=project_id,
        details={"format": "xlsx"},
    )

    requirements = db.scalars(
        select(Requirement)
        .where(Requirement.project_id == project_id)
        .order_by(Requirement.source_page.asc(), Requirement.created_at.asc())
    ).all()

    import io

    import xlsxwriter
    from fastapi.responses import StreamingResponse

    output = io.BytesIO()
    workbook = xlsxwriter.Workbook(output)
    worksheet = workbook.add_worksheet("Compliance Matrix")

    header_format = workbook.add_format(
        {"bold": True, "bg_color": "#4F46E5", "font_color": "#FFFFFF", "border": 1}
    )

    headers = [
        "Requirement ID",
        "Source Section",
        "Source Page",
        "Requirement Text",
        "Type",
        "Mandatory",
        "Status",
        "Owner",
        "Proposal Section",
        "Risk Level",
    ]

    for col_num, header in enumerate(headers):
        worksheet.write(0, col_num, header, header_format)

    for row_num, req in enumerate(requirements, start=1):
        worksheet.write(row_num, 0, str(req.id))
        worksheet.write(row_num, 1, req.source_section or "")
        worksheet.write(row_num, 2, req.source_page or "")
        worksheet.write(row_num, 3, req.original_text)
        worksheet.write(row_num, 4, req.requirement_type or "")
        worksheet.write(row_num, 5, "YES" if req.mandatory else "NO")
        worksheet.write(row_num, 6, req.status)
        worksheet.write(row_num, 7, req.owner_name or "")
        worksheet.write(row_num, 8, req.proposal_section or "")
        worksheet.write(row_num, 9, req.risk_level or "")

    workbook.close()
    output.seek(0)

    filename = f"compliance_matrix_{project_id}.xlsx"
    headers_resp = {"Content-Disposition": f'attachment; filename="{filename}"'}
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers=headers_resp,
    )


@router.get("/projects/{project_id}/export/proposal")
def export_proposal_docx(
    project_id: uuid.UUID,
    request: Request,
    db: Session = Depends(get_db),
) -> Any:
    org_id, user_id = get_current_org_and_user(request, db)
    project = get_project_for_org(db, project_id, org_id)

    log_audit_event(
        db,
        org_id=org_id,
        user_id=user_id,
        action="EXPORT_PROPOSAL_DOCX",
        entity_type="Project",
        entity_id=project_id,
        details={"format": "docx"},
    )

    from app.models.evidence import EvidenceLink
    from app.models.response import DraftResponse

    requirements = db.scalars(
        select(Requirement)
        .where(Requirement.project_id == project_id)
        .order_by(Requirement.source_page.asc(), Requirement.created_at.asc())
    ).all()

    import io
    from datetime import UTC, datetime

    from docx import Document as DocxDocument
    from fastapi.responses import StreamingResponse

    doc = DocxDocument()
    doc.add_heading(f"Proposal Response Draft: {project.name}", level=0)
    doc.add_paragraph(f"Client: {project.client_name}")
    doc.add_paragraph(f"Date generated: {datetime.now(UTC).strftime('%Y-%m-%d')}")
    doc.add_page_break()  # type: ignore[no-untyped-call]

    has_approved = False
    for req in requirements:
        draft = db.scalars(
            select(DraftResponse)
            .where(
                DraftResponse.requirement_id == req.id,
            )
            .order_by(DraftResponse.version.desc())
        ).first()

        is_approved = draft and draft.status == "approved"

        if is_approved:
            has_approved = True
            doc.add_heading(f"Requirement: {req.source_section or 'General'}", level=1)
            doc.add_paragraph(req.original_text, style="Normal")

            doc.add_heading("Answer Response", level=2)
            doc.add_paragraph(draft.content)  # type: ignore[union-attr]

            # Citation provenance — list all validated evidence links
            evidence_links = db.scalars(
                select(EvidenceLink).where(EvidenceLink.requirement_id == req.id)
            ).all()
            if evidence_links:
                doc.add_heading("Evidence Sources", level=3)
                for ev in evidence_links:
                    # Resolve document name for provenance
                    from app.models.document import Document as DocModel

                    ev_doc = db.scalar(
                        select(DocModel).where(DocModel.id == ev.document_id)
                    )
                    doc_name = ev_doc.name if ev_doc else str(ev.document_id)
                    page_ref = f", Page {ev.page_number}" if ev.page_number else ""
                    citation = f"[Source: {doc_name}{page_ref}]"
                    snippet_text = (
                        ev.snippet[:200] + "..."
                        if len(ev.snippet) > 200
                        else ev.snippet
                    )
                    doc.add_paragraph(f"{citation} {snippet_text}")
            doc.add_paragraph("")
        else:
            # Mark clearly in export
            status_label = (
                f" [{req.status}]"
                if req.status
                in ("NEEDS_REVIEW", "CHANGES_REQUESTED", "REJECTED", "NEEDS_EVIDENCE")
                else " [NOT STARTED]"
            )
            doc.add_heading(
                f"Requirement: {req.source_section or 'General'}{status_label}", level=1
            )
            doc.add_paragraph(req.original_text, style="Normal")

            doc.add_heading("Answer Response (Not Approved)", level=2)
            doc.add_paragraph("No response approved yet.")
            doc.add_paragraph("")

    if not has_approved:
        doc.add_paragraph("No approved answers available for export.")

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    filename = f"proposal_draft_{project_id}.docx"
    headers_resp = {"Content-Disposition": f'attachment; filename="{filename}"'}
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers_resp,
    )
