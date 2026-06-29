from pathlib import Path
from typing import Any

import docx
import fitz  # PyMuPDF
from fastapi import HTTPException, UploadFile

ALLOWED_EXTENSIONS = {".pdf", ".docx"}
ALLOWED_MIMETYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def validate_uploaded_file(file: UploadFile, max_size: int) -> None:
    """Validate extension, MIME type, empty file, and maximum size."""
    # Read up to max_size + 1 to determine if file is too large or empty
    content = file.file.read(max_size + 1)
    file.file.seek(0)

    if not content:
        raise HTTPException(status_code=400, detail="File is empty")

    if len(content) > max_size:
        raise HTTPException(status_code=400, detail="File exceeds maximum size limit")

    filename = file.filename or ""
    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400, detail=f"Unsupported file extension: {ext}"
        )

    mime_type = file.content_type
    if mime_type not in ALLOWED_MIMETYPES:
        raise HTTPException(
            status_code=400, detail=f"Unsupported MIME type: {mime_type}"
        )


def extract_pages(file_path: Path, file_type: str) -> list[dict[str, Any]]:
    """Extract pages from PDF or DOCX file.

    Returns:
        List of dicts: [{"page_number": int, "content": str}]
    """
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    # Determine extraction based on type or extension
    ext = file_path.suffix.lower()

    if ext == ".pdf" or file_type == "application/pdf":
        return _extract_pdf(file_path)
    elif (
        ext == ".docx"
        or file_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        return _extract_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type for extraction: {ext}")


def _extract_pdf(file_path: Path) -> list[dict[str, Any]]:
    pages = []
    doc = fitz.open(str(file_path))
    try:
        for idx, page in enumerate(doc):
            text = page.get_text()
            pages.append({"page_number": idx + 1, "content": text})
    finally:
        doc.close()
    return pages


def _extract_docx(file_path: Path) -> list[dict[str, Any]]:
    doc = docx.Document(str(file_path))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                parts.append(" | ".join(row_texts))

    chunk_words: list[str] = []
    chunks: list[str] = []
    for part in parts:
        words = part.split()
        chunk_words.extend(words)
        if len(chunk_words) >= 500:
            chunks.append(" ".join(chunk_words))
            chunk_words = []
    if chunk_words:
        chunks.append(" ".join(chunk_words))

    if not chunks:
        return [{"page_number": 1, "content": ""}]
    return [{"page_number": i + 1, "content": text} for i, text in enumerate(chunks)]
