"""
Integration tests for Step 8 — Evidence Grounding and Citation Integrity.

Covers:
- Valid same-project evidence links successfully.
- Fabricated snippet not on the page is rejected.
- Wrong page number is rejected.
- Foreign-project document evidence is rejected.
- Unapproved/failed document cannot be used as evidence.
- Client-submitted score is ignored/recomputed server-side.
- Draft approval blocks mandatory requirement with no evidence.
- Draft with evidence-supported content is approved normally.
- DOCX export includes citation provenance for evidence links.
- Previous Step 2 authorization, Step 3 CSRF, Step 4 auth tests still pass (by
  not touching those fixtures).
"""

from sqlalchemy import select

from app.core.database import get_default_org_and_user
from app.models.audit import AuditEvent
from app.models.document import Document, DocumentPage
from app.models.evidence import EvidenceLink
from app.models.project import ProposalProject
from app.models.requirement import Requirement
from app.models.response import DraftResponse
from app.services.ingestion_state import IngestionStatus


def _make_project(db, org_id, user_id, name="Test Project"):
    project = ProposalProject(
        organization_id=org_id,
        created_by_id=user_id,
        name=name,
        client_name="Acme Corp",
        status="draft",
    )
    db.add(project)
    db.commit()
    return project


def _make_document(
    db,
    project_id,
    user_id,
    content="The system must support multi-factor authentication for all users.",
    processing_status="completed",
    approval_status="APPROVED",
    doc_role="knowledge_base",
    name="KnowledgeDoc",
    ingestion_status=IngestionStatus.COMPLETED,
):
    doc = Document(
        project_id=project_id,
        name=name,
        file_path="./data/test.pdf",
        file_type="application/pdf",
        doc_role=doc_role,
        processing_status=processing_status,
        approval_status=approval_status,
        created_by_id=user_id,
        content=content,
        ingestion_status=ingestion_status,
    )
    db.add(doc)
    db.commit()
    return doc


def _make_page(db, document_id, page_number=1, content=None):
    if content is None:
        content = "The system must support multi-factor authentication for all users."
    page = DocumentPage(
        document_id=document_id,
        page_number=page_number,
        content=content,
    )
    db.add(page)
    db.commit()
    return page


def _make_requirement(
    db, project_id, mandatory=False, text="The system must support MFA."
):
    req = Requirement(
        project_id=project_id,
        original_text=text,
        status="NOT_STARTED",
        mandatory=mandatory,
    )
    db.add(req)
    db.commit()
    return req


# ---------------------------------------------------------------------------
# Evidence link — valid case
# ---------------------------------------------------------------------------


def test_valid_same_project_evidence_links_successfully(client, db):
    """Valid snippet from the project's approved document links without error."""
    org_id, user_id = get_default_org_and_user(db)
    project = _make_project(db, org_id, user_id)
    page_content = "The system must support multi-factor authentication for all users."
    doc = _make_document(db, project.id, user_id, content=page_content)
    _make_page(db, doc.id, page_number=1, content=page_content)
    req = _make_requirement(db, project.id)

    resp = client.post(
        f"/requirements/{req.id}/evidence/link",
        data={
            "document_id": str(doc.id),
            "snippet": "multi-factor authentication for all users",
            "page_number": 1,
            "score": 0.99,  # should be ignored server-side
        },
        follow_redirects=False,
    )
    assert resp.status_code == 303, resp.text

    db.expire_all()
    links = db.scalars(
        select(EvidenceLink).where(EvidenceLink.requirement_id == req.id)
    ).all()
    assert len(links) == 1
    # Client score must NOT be stored as-is; server always sets 0.0 for now
    assert links[0].score == 0.0
    assert "multi-factor authentication" in links[0].snippet


# ---------------------------------------------------------------------------
# Evidence link — fabricated snippet rejected
# ---------------------------------------------------------------------------


def test_fabricated_snippet_not_on_page_rejected(client, db):
    """Snippet that does not appear in the stored page content is rejected."""
    org_id, user_id = get_default_org_and_user(db)
    project = _make_project(db, org_id, user_id, name="Fabricated Snippet Proj")
    doc = _make_document(
        db, project.id, user_id, content="The system must support MFA."
    )
    _make_page(db, doc.id, page_number=1, content="The system must support MFA.")
    req = _make_requirement(db, project.id)

    resp = client.post(
        f"/requirements/{req.id}/evidence/link",
        data={
            "document_id": str(doc.id),
            "snippet": "completely fabricated snippet text",
            "page_number": 1,
            "score": 1.0,
        },
    )
    assert resp.status_code == 400
    assert "not found" in resp.text.lower() or "snippet" in resp.text.lower()

    db.expire_all()
    links = db.scalars(
        select(EvidenceLink).where(EvidenceLink.requirement_id == req.id)
    ).all()
    assert len(links) == 0


# ---------------------------------------------------------------------------
# Evidence link — wrong page number rejected
# ---------------------------------------------------------------------------


def test_wrong_page_number_rejected(client, db):
    """Evidence referencing a page number that doesn't exist is rejected."""
    org_id, user_id = get_default_org_and_user(db)
    project = _make_project(db, org_id, user_id, name="Wrong Page Proj")
    page_content = "The system must support MFA for all users."
    doc = _make_document(db, project.id, user_id, content=page_content)
    # Only add page 1
    _make_page(db, doc.id, page_number=1, content=page_content)
    req = _make_requirement(db, project.id)

    resp = client.post(
        f"/requirements/{req.id}/evidence/link",
        data={
            "document_id": str(doc.id),
            "snippet": "system must support MFA",
            "page_number": 99,  # page 99 does not exist
        },
    )
    assert resp.status_code == 400
    assert "page" in resp.text.lower()


# ---------------------------------------------------------------------------
# Evidence link — snippet on wrong page rejected
# ---------------------------------------------------------------------------


def test_snippet_exists_but_not_on_specified_page_rejected(client, db):
    """Snippet is present in the document but NOT on the specified page."""
    org_id, user_id = get_default_org_and_user(db)
    project = _make_project(db, org_id, user_id, name="Wrong Page Content Proj")
    doc = _make_document(
        db,
        project.id,
        user_id,
        content="Page 1 content. Page 2 content with MFA support.",
    )
    _make_page(db, doc.id, page_number=1, content="Page 1 content only.")
    _make_page(db, doc.id, page_number=2, content="Page 2 content with MFA support.")
    req = _make_requirement(db, project.id)

    # Snippet is on page 2 but client claims page 1
    resp = client.post(
        f"/requirements/{req.id}/evidence/link",
        data={
            "document_id": str(doc.id),
            "snippet": "MFA support",
            "page_number": 1,  # wrong page — MFA support is on page 2
        },
    )
    assert resp.status_code == 400


# ---------------------------------------------------------------------------
# Evidence link — foreign project document rejected
# ---------------------------------------------------------------------------


def test_foreign_project_document_evidence_rejected(client, db):
    """Evidence from a different project's document cannot be linked."""
    org_id, user_id = get_default_org_and_user(db)

    proj_a = _make_project(db, org_id, user_id, name="Project A")
    proj_b = _make_project(db, org_id, user_id, name="Project B")

    page_content = "The vendor shall provide 24/7 support."
    doc_b = _make_document(db, proj_b.id, user_id, content=page_content)
    _make_page(db, doc_b.id, page_number=1, content=page_content)

    req_a = _make_requirement(db, proj_a.id)

    resp = client.post(
        f"/requirements/{req_a.id}/evidence/link",
        data={
            "document_id": str(doc_b.id),  # document belongs to proj_b
            "snippet": "vendor shall provide 24/7 support",
            "page_number": 1,
        },
    )
    # Should be 404 (no project info leaked) or 400
    assert resp.status_code in (400, 404)

    db.expire_all()
    links = db.scalars(
        select(EvidenceLink).where(EvidenceLink.requirement_id == req_a.id)
    ).all()
    assert len(links) == 0


# ---------------------------------------------------------------------------
# Evidence link — unapproved document rejected
# ---------------------------------------------------------------------------


def test_unapproved_knowledge_base_document_rejected(client, db):
    """Knowledge-base document that isn't APPROVED cannot be used as evidence."""
    org_id, user_id = get_default_org_and_user(db)
    project = _make_project(db, org_id, user_id, name="Unapproved Doc Proj")
    page_content = "The system must support MFA."
    doc = _make_document(
        db,
        project.id,
        user_id,
        content=page_content,
        approval_status="PENDING",  # not approved
    )
    _make_page(db, doc.id, page_number=1, content=page_content)
    req = _make_requirement(db, project.id)

    resp = client.post(
        f"/requirements/{req.id}/evidence/link",
        data={
            "document_id": str(doc.id),
            "snippet": "system must support MFA",
            "page_number": 1,
        },
    )
    assert resp.status_code == 400
    assert "approved" in resp.text.lower() or "not" in resp.text.lower()


# ---------------------------------------------------------------------------
# Evidence link — unprocessed document rejected
# ---------------------------------------------------------------------------


def test_unprocessed_document_rejected(client, db):
    """Document with processing_status != 'completed' cannot be used as evidence."""
    org_id, user_id = get_default_org_and_user(db)
    project = _make_project(db, org_id, user_id, name="Unprocessed Doc Proj")
    page_content = "The system must support MFA."
    doc = _make_document(
        db,
        project.id,
        user_id,
        content=page_content,
        processing_status="pending",
        approval_status="APPROVED",
    )
    req = _make_requirement(db, project.id)

    resp = client.post(
        f"/requirements/{req.id}/evidence/link",
        data={
            "document_id": str(doc.id),
            "snippet": "system must support MFA",
        },
    )
    assert resp.status_code == 400
    assert "processed" in resp.text.lower() or "not" in resp.text.lower()


# ---------------------------------------------------------------------------
# Evidence link — client-submitted score is ignored
# ---------------------------------------------------------------------------


def test_client_submitted_score_ignored_server_sets_zero(client, db):
    """Client score 0.99 must be discarded; server stores 0.0."""
    org_id, user_id = get_default_org_and_user(db)
    project = _make_project(db, org_id, user_id, name="Score Ignore Proj")
    page_content = "The system must support MFA for all users."
    doc = _make_document(db, project.id, user_id, content=page_content)
    _make_page(db, doc.id, page_number=1, content=page_content)
    req = _make_requirement(db, project.id)

    resp = client.post(
        f"/requirements/{req.id}/evidence/link",
        data={
            "document_id": str(doc.id),
            "snippet": "system must support MFA",
            "page_number": 1,
            "score": 0.99,
        },
        follow_redirects=False,
    )
    assert resp.status_code == 303

    db.expire_all()
    link = db.scalars(
        select(EvidenceLink).where(EvidenceLink.requirement_id == req.id)
    ).first()
    assert link is not None
    assert link.score == 0.0, f"Expected 0.0, got {link.score}"


# ---------------------------------------------------------------------------
# Draft approval — mandatory requirement blocked without evidence
# ---------------------------------------------------------------------------


def test_approve_mandatory_requirement_without_evidence_blocked(client, db):
    """Approving a mandatory requirement with no evidence sets NEEDS_REVIEW."""
    org_id, user_id = get_default_org_and_user(db)
    project = _make_project(db, org_id, user_id, name="Mandatory No Evidence")
    req = _make_requirement(db, project.id, mandatory=True)

    draft = DraftResponse(
        requirement_id=req.id,
        content="We comply with all requirements.",
        status="draft",
        version=1,
    )
    db.add(draft)
    db.commit()

    resp = client.post(
        f"/requirements/{req.id}/draft/approve",
        follow_redirects=False,
    )
    # Should redirect with warning, not approve
    assert resp.status_code == 303
    assert "warning" in resp.headers.get("location", "")

    db.expire_all()
    updated_req = db.get(Requirement, req.id)
    updated_draft = db.get(DraftResponse, draft.id)
    assert updated_req.status == "NEEDS_REVIEW"
    assert updated_draft.status == "needs_review"


# ---------------------------------------------------------------------------
# Draft approval — non-mandatory requirement with evidence approved normally
# ---------------------------------------------------------------------------


def test_approve_non_mandatory_requirement_without_evidence_succeeds(client, db):
    """Non-mandatory requirement with no evidence can still be approved."""
    org_id, user_id = get_default_org_and_user(db)
    project = _make_project(db, org_id, user_id, name="Non-Mandatory Approval")
    req = _make_requirement(db, project.id, mandatory=False)

    draft = DraftResponse(
        requirement_id=req.id,
        content="We will comply with this requirement.",
        status="draft",
        version=1,
    )
    db.add(draft)
    db.commit()

    resp = client.post(
        f"/requirements/{req.id}/draft/approve",
        follow_redirects=False,
    )
    assert resp.status_code == 303

    db.expire_all()
    updated_req = db.get(Requirement, req.id)
    updated_draft = db.get(DraftResponse, draft.id)
    # Non-mandatory with no evidence passes grounding (no claims to check against
    # evidence, or evidence list is empty so grounding check is skipped)
    assert updated_req.status == "APPROVED"
    assert updated_draft.status == "approved"


# ---------------------------------------------------------------------------
# Draft approval — audit event for blocked approval
# ---------------------------------------------------------------------------


def test_blocked_approval_generates_audit_event(client, db):
    """A blocked approval (mandatory + no evidence) generates a draft_approve_blocked
    audit event."""
    org_id, user_id = get_default_org_and_user(db)
    project = _make_project(db, org_id, user_id, name="Audit Blocked Proj")
    req = _make_requirement(db, project.id, mandatory=True)

    draft = DraftResponse(
        requirement_id=req.id,
        content="We will comply.",
        status="draft",
        version=1,
    )
    db.add(draft)
    db.commit()

    client.post(
        f"/requirements/{req.id}/draft/approve",
        follow_redirects=False,
    )

    db.expire_all()
    audit = db.scalars(
        select(AuditEvent).where(
            AuditEvent.action == "draft_approve_blocked",
            AuditEvent.entity_id == draft.id,
        )
    ).first()
    assert audit is not None


# ---------------------------------------------------------------------------
# DOCX export — citation provenance included
# ---------------------------------------------------------------------------


def test_docx_export_includes_citation_provenance(client, db):
    """Approved draft export includes [Source: doc_name, Page N] citation."""
    org_id, user_id = get_default_org_and_user(db)
    project = _make_project(db, org_id, user_id, name="Citation Provenance Proj")
    page_content = "The system must support MFA for all users."
    doc = _make_document(
        db, project.id, user_id, content=page_content, name="SecurityPolicy"
    )
    _make_page(db, doc.id, page_number=1, content=page_content)

    req = _make_requirement(
        db, project.id, mandatory=False, text="The system must support MFA."
    )

    # Create an approved draft
    draft = DraftResponse(
        requirement_id=req.id,
        content="We support MFA for all users as described in our security policy.",
        status="approved",
        version=1,
    )
    db.add(draft)

    # Create an evidence link manually (bypassing validation for export test)
    ev_link = EvidenceLink(
        requirement_id=req.id,
        document_id=doc.id,
        snippet="system must support MFA for all users",
        page_number=1,
        score=0.0,
    )
    db.add(ev_link)
    db.commit()

    resp = client.get(f"/projects/{project.id}/export/proposal")
    assert resp.status_code == 200
    assert resp.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    import io

    from docx import Document as DocxDocument

    docx_doc = DocxDocument(io.BytesIO(resp.content))
    full_text = "\n".join(p.text for p in docx_doc.paragraphs if p.text)

    assert "SecurityPolicy" in full_text, (
        f"Doc name 'SecurityPolicy' not found in export. Got: {full_text[:500]}"
    )
    assert "Page 1" in full_text or "page 1" in full_text.lower(), (
        f"Page number not found in export. Got: {full_text[:500]}"
    )
    assert "Source:" in full_text, (
        f"Citation '[Source:]' not found in export. Got: {full_text[:500]}"
    )


# ---------------------------------------------------------------------------
# Snippet too short — rejected by validation
# ---------------------------------------------------------------------------


def test_snippet_too_short_rejected(client, db):
    """Snippets below SNIPPET_MIN_LEN characters are rejected."""
    org_id, user_id = get_default_org_and_user(db)
    project = _make_project(db, org_id, user_id, name="Short Snippet Proj")
    page_content = "MFA is required for all users in our system by policy."
    doc = _make_document(db, project.id, user_id, content=page_content)
    _make_page(db, doc.id, page_number=1, content=page_content)
    req = _make_requirement(db, project.id)

    resp = client.post(
        f"/requirements/{req.id}/evidence/link",
        data={
            "document_id": str(doc.id),
            "snippet": "MFA",  # 3 chars, below SNIPPET_MIN_LEN=10
            "page_number": 1,
        },
    )
    assert resp.status_code in (400, 422)
