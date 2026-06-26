import io
from pathlib import Path

import docx
import pytest
from fastapi import HTTPException, UploadFile
from reportlab.pdfgen import canvas

from app.services.extractor import (
    extract_pages,
    validate_uploaded_file,
)


def create_synthetic_pdf(dest_path: Path, pages_text: list[str]) -> None:
    c = canvas.Canvas(str(dest_path))
    for text in pages_text:
        c.drawString(100, 750, text)
        c.showPage()
    c.save()


def create_synthetic_docx(dest_path: Path, text_content: str) -> None:
    doc = docx.Document()
    doc.add_paragraph(text_content)
    # add a simple table to test table extraction
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Header A"
    hdr_cells[1].text = "Header B"
    doc.save(str(dest_path))


def test_validate_uploaded_file_valid_pdf():
    file = UploadFile(
        filename="rfp.pdf",
        file=io.BytesIO(b"%PDF-1.4 mock content"),
        headers={"content-type": "application/pdf"},
    )
    # Should not raise exception
    validate_uploaded_file(file, max_size=1024)


def test_validate_uploaded_file_valid_docx():
    file = UploadFile(
        filename="rfp.docx",
        file=io.BytesIO(b"mock docx content"),
        headers={
            "content-type": (
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document"
            )
        },
    )
    # Should not raise exception
    validate_uploaded_file(file, max_size=1024)


def test_validate_uploaded_file_empty():
    file = UploadFile(
        filename="empty.pdf",
        file=io.BytesIO(b""),
        headers={"content-type": "application/pdf"},
    )
    with pytest.raises(HTTPException) as exc:
        validate_uploaded_file(file, max_size=1024)
    assert exc.value.status_code == 400
    assert "empty" in exc.value.detail.lower()


def test_validate_uploaded_file_too_large():
    file = UploadFile(
        filename="large.pdf",
        file=io.BytesIO(b"a" * 100),
        headers={"content-type": "application/pdf"},
    )
    with pytest.raises(HTTPException) as exc:
        validate_uploaded_file(file, max_size=50)
    assert exc.value.status_code == 400
    assert "exceeds" in exc.value.detail.lower()


def test_validate_uploaded_file_invalid_ext():
    file = UploadFile(
        filename="test.txt",
        file=io.BytesIO(b"some text"),
        headers={"content-type": "text/plain"},
    )
    with pytest.raises(HTTPException) as exc:
        validate_uploaded_file(file, max_size=1024)
    assert exc.value.status_code == 400
    assert "extension" in exc.value.detail.lower()


def test_validate_uploaded_file_invalid_mime():
    file = UploadFile(
        filename="test.pdf",
        file=io.BytesIO(b"some text"),
        headers={"content-type": "text/plain"},
    )
    with pytest.raises(HTTPException) as exc:
        validate_uploaded_file(file, max_size=1024)
    assert exc.value.status_code == 400
    assert "mime" in exc.value.detail.lower()


def test_extract_pdf_pages(tmp_path: Path):
    pdf_path = tmp_path / "test.pdf"
    create_synthetic_pdf(pdf_path, ["Page One Text", "Page Two Text"])

    pages = extract_pages(pdf_path, "application/pdf")
    assert len(pages) == 2
    assert pages[0]["page_number"] == 1
    assert "Page One" in pages[0]["content"]
    assert pages[1]["page_number"] == 2
    assert "Page Two" in pages[1]["content"]


def test_extract_docx_pages(tmp_path: Path):
    docx_path = tmp_path / "test.docx"
    create_synthetic_docx(docx_path, "This is paragraph text.")

    pages = extract_pages(
        docx_path,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert len(pages) == 1
    assert pages[0]["page_number"] == 1
    assert "This is paragraph text" in pages[0]["content"]
    assert "Header A | Header B" in pages[0]["content"]
