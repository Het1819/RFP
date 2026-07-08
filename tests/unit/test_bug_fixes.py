from unittest.mock import AsyncMock, MagicMock, patch

import docx
import pytest
from sqlalchemy import select

from app.core.database import get_default_org_and_user
from app.core.llm import AnthropicProvider, FakeLLMProvider
from app.models.document import Document, DocumentPage
from app.models.project import ProposalProject
from app.models.requirement import Requirement
from app.models.response import DraftResponse
from app.services.extractor import _extract_docx
from app.services.project_service import process_document_background
from app.services.retriever import retrieve_evidence


def test_process_document_background_skips_run_extraction_sync(db):
    org_id, user_id = get_default_org_and_user(db)
    project = ProposalProject(
        organization_id=org_id,
        created_by_id=user_id,
        name="Skipping Extraction Proj",
        client_name="Client A",
        status="draft",
    )
    db.add(project)
    db.commit()

    doc = Document(
        project_id=project.id,
        created_by_id=user_id,
        name="kb.pdf",
        file_path="/tmp/kb.pdf",
        file_type="application/pdf",
        doc_role="knowledge_base",
    )
    db.add(doc)
    db.commit()

    with (
        patch("app.services.project_service.extract_pages") as mock_extract,
        patch("app.services.extraction_service.run_extraction_sync") as mock_run_ext,
    ):
        mock_extract.return_value = [{"page_number": 1, "content": "Knowledge content"}]

        from app.core.database import SessionLocal

        process_document_background(SessionLocal, doc.id)

        mock_run_ext.assert_not_called()


@pytest.mark.asyncio
async def test_anthropic_extract_requirements_parses_page_markers():
    mock_client = MagicMock()
    mock_messages = AsyncMock()
    mock_response = MagicMock()
    mock_content = MagicMock()
    mock_content.text = (
        '[{"original_text": "Must do X", "source_section": "1.2", '
        '"source_page": 2, "requirement_type": "Technical", '
        '"mandatory": true, "risk_level": "Low"}]'
    )
    mock_response.content = [mock_content]
    mock_messages.create.return_value = mock_response
    mock_client.messages = mock_messages

    provider = AnthropicProvider(api_key="fake-key", model="claude-sonnet-4-6")
    provider.client = mock_client

    text_input = "[PAGE 2]\nMust do X"
    res = await provider.extract_requirements(text_input)

    assert len(res) == 1
    assert res[0].source_page == 2

    kwargs = mock_messages.create.call_args.kwargs
    assert "[PAGE 2]" in kwargs["messages"][0]["content"]


@pytest.mark.asyncio
async def test_fake_llm_provider_classifies_types():
    provider = FakeLLMProvider()

    res_comp = await provider.extract_requirements(
        "You must have a certif of operation."
    )
    assert len(res_comp) == 1
    assert res_comp[0].requirement_type == "Compliance"

    res_comm = await provider.extract_requirements(
        "You must pay a fixed interest rate."
    )
    assert len(res_comm) == 1
    assert res_comm[0].requirement_type == "Commercial"

    res_proc = await provider.extract_requirements(
        "You must submit before the deadline."
    )
    assert len(res_proc) == 1
    assert res_proc[0].requirement_type == "Procedural"


def test_extract_docx_exceeds_500_words(tmp_path):
    doc_path = tmp_path / "long_doc.docx"
    doc = docx.Document()
    doc.add_paragraph("word " * 500)
    doc.add_paragraph("word " * 10)
    doc.save(str(doc_path))

    pages = _extract_docx(doc_path)
    assert len(pages) == 2
    assert pages[0]["page_number"] == 1
    assert pages[1]["page_number"] == 2


def test_retrieve_evidence_preserves_hyphens(db):
    org_id, user_id = get_default_org_and_user(db)

    project = ProposalProject(
        organization_id=org_id,
        created_by_id=user_id,
        name="Hyphen Proj",
        client_name="Client H",
        status="draft",
    )
    db.add(project)
    db.commit()

    doc = Document(
        project_id=project.id,
        created_by_id=user_id,
        name="kb.pdf",
        file_path="/tmp/kb.pdf",
        file_type="application/pdf",
        doc_role="knowledge_base",
        processing_status="completed",
        approval_status="APPROVED",
    )
    db.add(doc)
    db.commit()

    page = DocumentPage(
        document_id=doc.id,
        page_number=1,
        content="We provide AD Category-I classification support.",
    )
    db.add(page)
    db.commit()

    res = retrieve_evidence(db, project.id, "AD Category-I")
    assert len(res) == 1
    assert "AD Category-I" in res[0]["snippet"]


def test_update_requirement_action_rejects_invalid_status(client, db):
    org_id, user_id = get_default_org_and_user(db)
    project = ProposalProject(
        organization_id=org_id,
        created_by_id=user_id,
        name="Test Validation Proj",
        client_name="Client V",
        status="draft",
    )
    db.add(project)
    db.commit()

    req = Requirement(
        project_id=project.id,
        original_text="Requirement text",
        status="NOT_STARTED",
    )
    db.add(req)
    db.commit()

    payload = {
        "original_text": "Updated Text",
        "status": "INVALID_STATUS_XYZ",
    }
    response = client.post(f"/requirements/{req.id}/edit", data=payload)
    assert response.status_code == 422


@pytest.mark.asyncio
async def test_draft_generation_creates_new_version(client, db):
    org_id, user_id = get_default_org_and_user(db)
    project = ProposalProject(
        organization_id=org_id,
        created_by_id=user_id,
        name="Draft Versioning Proj",
        client_name="Client D",
        status="draft",
    )
    db.add(project)
    db.commit()

    req = Requirement(
        project_id=project.id,
        original_text="Need draft versioning test.",
        status="NOT_STARTED",
    )
    db.add(req)
    db.commit()

    from app.core.llm import DraftResponseDraft

    mock_draft_response = DraftResponseDraft(
        answer_text="Mock answer text",
        confidence=0.9,
        needs_evidence=False,
        assumptions=None,
    )

    with patch(
        "app.core.llm.FakeLLMProvider.draft_response", new_callable=AsyncMock
    ) as mock_draft:
        mock_draft.return_value = mock_draft_response

        resp1 = client.post(f"/requirements/{req.id}/draft", follow_redirects=False)
        assert resp1.status_code == 303

        resp2 = client.post(f"/requirements/{req.id}/draft", follow_redirects=False)
        assert resp2.status_code == 303

    db.expire_all()
    drafts = db.scalars(
        select(DraftResponse)
        .where(DraftResponse.requirement_id == req.id)
        .order_by(DraftResponse.version.asc())
    ).all()

    assert len(drafts) == 2
    assert drafts[0].version == 1
    assert drafts[1].version == 2
