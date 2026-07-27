"""Unit test suite for isolated document parser service."""

import hashlib
import io
from pathlib import Path

import docx
import fitz
import pytest
from fastapi.testclient import TestClient

from app.parser_service.config import (
    DOCX_MIME,
    PDF_MIME,
)
from app.parser_service.docx_extractor import extract_docx_units
from app.parser_service.main import app
from app.parser_service.pdf_extractor import extract_pdf_units


@pytest.fixture
def client():
    return TestClient(app)


def _create_sample_pdf(pages_text: list[str]) -> bytes:
    doc = fitz.open()
    for text in pages_text:
        page = doc.new_page()
        page.insert_text((50, 50), text)
    pdf_bytes = doc.tobytes()
    doc.close()
    return pdf_bytes


def _create_sample_docx(
    paragraphs: list[str], table_rows: list[list[str]] | None = None
) -> bytes:
    doc = docx.Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    if table_rows:
        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r_idx, row in enumerate(table_rows):
            for c_idx, val in enumerate(row):
                table.cell(r_idx, c_idx).text = val
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


class TestParserServiceEndpoints:
    def test_healthz_endpoint(self, client: TestClient):
        response = client.get("/healthz")
        assert response.status_code == 200
        data = response.json()
        assert data["status"] == "ok"
        assert data["service"] == "rfp-isolated-parser"
        assert "version" in data

    def test_parse_clean_pdf_success(self, client: TestClient):
        pdf_bytes = _create_sample_pdf(["Page 1 Content", "Page 2 Content"])
        sha256 = hashlib.sha256(pdf_bytes).hexdigest()
        size = len(pdf_bytes)

        response = client.post(
            "/parse",
            headers={
                "X-Content-Type": PDF_MIME,
                "X-Expected-SHA256": sha256,
                "X-Expected-Size-Bytes": str(size),
            },
            files={"file": ("test.pdf", pdf_bytes, PDF_MIME)},
        )

        assert response.status_code == 200
        data = response.json()
        assert data["protocol_version"] == "1.0"
        assert data["document_type"] == "PDF"
        assert data["total_units"] == 2
        assert len(data["units"]) == 2
        assert data["units"][0]["unit_kind"] == "PDF_PAGE"
        assert data["units"][0]["source_locator"] == "page_1"
        assert "Page 1 Content" in data["units"][0]["content"]
        assert data["units"][1]["source_locator"] == "page_2"
        assert "Page 2 Content" in data["units"][1]["content"]

    def test_parse_clean_docx_success(self, client: TestClient):
        docx_bytes = _create_sample_docx(
            ["Paragraph 1 text", "Paragraph 2 text"],
            table_rows=[["Header 1", "Header 2"], ["Val 1", "Val 2"]],
        )
        sha256 = hashlib.sha256(docx_bytes).hexdigest()
        size = len(docx_bytes)

        response = client.post(
            "/parse",
            headers={
                "X-Content-Type": DOCX_MIME,
                "X-Expected-SHA256": sha256,
                "X-Expected-Size-Bytes": str(size),
            },
            files={"file": ("test.docx", docx_bytes, DOCX_MIME)},
        )

        assert response.status_code == 200
        data = response.json()
        assert data["protocol_version"] == "1.0"
        assert data["document_type"] == "DOCX"
        assert data["total_units"] >= 1
        assert data["units"][0]["unit_kind"] == "DOCX_LOGICAL_CHUNK"
        assert data["units"][0]["source_locator"].startswith("chunk_")
        assert "Paragraph 1 text" in data["units"][0]["content"]
        assert "Header 1 | Header 2" in data["units"][0]["content"]

    def test_unsupported_content_type_rejected(self, client: TestClient):
        response = client.post(
            "/parse",
            headers={
                "X-Content-Type": "image/png",
                "X-Expected-SHA256": "a" * 64,
                "X-Expected-Size-Bytes": "100",
            },
            files={"file": ("image.png", b"fake png data", "image/png")},
        )
        assert response.status_code == 400
        data = response.json()
        assert data["error_code"] == "UNSUPPORTED_CONTENT_TYPE"

    def test_digest_mismatch_rejected(self, client: TestClient):
        pdf_bytes = _create_sample_pdf(["Content"])
        size = len(pdf_bytes)
        wrong_sha256 = "0" * 64

        response = client.post(
            "/parse",
            headers={
                "X-Content-Type": PDF_MIME,
                "X-Expected-SHA256": wrong_sha256,
                "X-Expected-Size-Bytes": str(size),
            },
            files={"file": ("test.pdf", pdf_bytes, PDF_MIME)},
        )
        assert response.status_code == 400
        data = response.json()
        assert data["error_code"] == "DIGEST_MISMATCH"

    def test_size_mismatch_rejected(self, client: TestClient):
        pdf_bytes = _create_sample_pdf(["Content"])
        sha256 = hashlib.sha256(pdf_bytes).hexdigest()
        wrong_size = len(pdf_bytes) + 100

        response = client.post(
            "/parse",
            headers={
                "X-Content-Type": PDF_MIME,
                "X-Expected-SHA256": sha256,
                "X-Expected-Size-Bytes": str(wrong_size),
            },
            files={"file": ("test.pdf", pdf_bytes, PDF_MIME)},
        )
        assert response.status_code == 400
        data = response.json()
        assert data["error_code"] == "SIZE_MISMATCH"

    def test_malformed_pdf_fails_closed(self, client: TestClient):
        corrupt_bytes = b"%PDF-1.4 corrupt data stream without trailer"
        sha256 = hashlib.sha256(corrupt_bytes).hexdigest()
        size = len(corrupt_bytes)

        response = client.post(
            "/parse",
            headers={
                "X-Content-Type": PDF_MIME,
                "X-Expected-SHA256": sha256,
                "X-Expected-Size-Bytes": str(size),
            },
            files={"file": ("corrupt.pdf", corrupt_bytes, PDF_MIME)},
        )
        assert response.status_code == 422
        data = response.json()
        assert data["error_code"] == "PDF_CORRUPT"

    def test_malformed_docx_fails_closed(self, client: TestClient):
        corrupt_bytes = b"PK\x03\x04 corrupt zip archive data"
        sha256 = hashlib.sha256(corrupt_bytes).hexdigest()
        size = len(corrupt_bytes)

        response = client.post(
            "/parse",
            headers={
                "X-Content-Type": DOCX_MIME,
                "X-Expected-SHA256": sha256,
                "X-Expected-Size-Bytes": str(size),
            },
            files={"file": ("corrupt.docx", corrupt_bytes, DOCX_MIME)},
        )
        assert response.status_code == 422
        data = response.json()
        assert data["error_code"] == "DOCX_CORRUPT"


class TestExtractorUnitsAndBoundaries:
    def test_pdf_page_limits(self, tmp_path: Path):
        pages = [f"Page {i}" for i in range(1, 10)]
        pdf_bytes = _create_sample_pdf(pages)
        pdf_file = tmp_path / "test.pdf"
        pdf_file.write_bytes(pdf_bytes)

        units, total_units, _total_chars = extract_pdf_units(pdf_file)
        assert total_units == 9
        assert len(units) == 9
        assert units[0].sequence == 1
        assert units[0].unit_kind == "PDF_PAGE"
        assert units[0].source_locator == "page_1"

    def test_docx_logical_chunks(self, tmp_path: Path):
        paragraphs = [f"Paragraph {i}: " + ("text " * 50) for i in range(10)]
        docx_bytes = _create_sample_docx(paragraphs)
        docx_file = tmp_path / "test.docx"
        docx_file.write_bytes(docx_bytes)

        units, total_units, _total_chars = extract_docx_units(docx_file)
        assert total_units >= 1
        assert units[0].unit_kind == "DOCX_LOGICAL_CHUNK"
        assert units[0].source_locator.startswith("chunk_")

    def test_text_normalization_strips_control_characters(self):
        from app.parser_service.normalizer import normalize_text

        raw_text = "Hello\x00 World!\x07\x1b\x08 Good\nText"
        norm = normalize_text(raw_text)
        assert "\x00" not in norm
        assert "\x07" not in norm
        assert norm == "Hello World! Good\nText"
